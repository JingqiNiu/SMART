#!/usr/bin/env python
# coding: utf-8

import os
import os.path as osp
import argparse
import subprocess

# pip install python-docx
from docx import Document
from docx.shared import Inches

from pd_helper import merge_pred_with_gt
from tools_utils import open_file_of_exp, data_distribution, get_csv_from_data_name


def add_table_to_docx(doc, df):
    """
    Inserts a table into the document

    Args:
        doc (docx.Document): 
        df (pd.DataFrame): 
    """
    row_num = df.shape[0] + 1
    col_num = df.shape[1] + 1

    df = df.T.reset_index().T
    df = df.reset_index()

    tb = doc.add_table(rows=row_num, cols=col_num)

    for row in range(row_num):
        row_cells = tb.rows[row].cells
        for col in range(col_num):
            row_cells[col].text = str(df.iloc[row, col])
    tb.rows[0].cells[0].text = ''


def generate_docx():
    document = Document()
    document.add_heading('训练报告', 0)

    document.add_heading('一、概述', level=1)

    document.add_heading('二、数据集', level=1)

    document.add_paragraph('训练集的分布：')
    add_table_to_docx(document, train_data_table)
    document.add_paragraph('验证集的分布：')
    add_table_to_docx(document, val_data_table)
    document.add_paragraph('测试集的分布：')
    add_table_to_docx(document, test_data_table)

    document.add_heading('三、模型和结果', level=1)

    document.add_heading('3.1 模型参数', level=2)

    document.add_heading('3.2 结果', level=2)

    if 'cls_report.txt' in train_result_dict:
        document.add_paragraph(['训练集的结果为：\n'])
        document.add_paragraph(train_result_dict['cls_report.txt'])

    document.add_paragraph(['验证集的结果为：\n'])
    document.add_paragraph(test_result_dict['cls_report.txt'])
    roc_path = os.path.join(exp_dir, 'test_result', 'roc.jpg')
    if os.path.exists(roc_path):
        document.add_picture(roc_path, width=Inches(3.5))

    document.add_heading('四、可视化图', level=1)

    document.add_heading('五、附录', level=1)

    document.add_paragraph('实验的参数为：\n')
    cfg_txt = str(cfg)
    cfg_txt = cfg_txt.replace(',', '\n')
    document.add_paragraph(cfg_txt)
    save_doc_name = os.path.join(exp_dir, 'training_report.docx')
    document.save(save_doc_name)
    print('save docx file: ', save_doc_name)
    return 


if __name__ == '__main__':

    parser = argparse.ArgumentParser(description='process after experiment.')
    parser.add_argument('exp_dir', type=str, default='',help='the path to experiment')

    parser.add_argument('-tb', '--tensorboard', action='store_true', default=False,
                        help='Whether or not use tensorboard')

    args = parser.parse_args()
    exp_dir = args.exp_dir

    # The file directory of the experiment output
    test_result_dir = osp.join(exp_dir, 'test_result/')
    train_result_dir = osp.join(exp_dir, 'evaluate_train_result/')
    ck_dir = osp.join(exp_dir, 'checkpoints/')

    # Read files
    exp_dict = open_file_of_exp(exp_dir)
    test_result_dict = open_file_of_exp(test_result_dir)
    train_result_dict = open_file_of_exp(train_result_dir)

    #  data_distribution
    cfg = exp_dict['hparams.yaml']
    label_name = cfg['label_name']
    dataset_define = cfg.get('dataset_define', None)

    train_data_table = data_distribution(cfg['train_data_name'], label_name, dataset_define)
    val_data_table = data_distribution(cfg['val_data_name'], label_name, dataset_define)
    test_data_table = data_distribution(cfg['test_data_name'], label_name, dataset_define)

    # print experimental information

    # generate_docx
    generate_docx()

    # merge the predicted CSV with GT's CSV
    pred_name = os.path.join(test_result_dir, 'report_table.csv')
    gt_csv = get_csv_from_data_name(cfg['test_data_name'], label_name)
    merge_csv = merge_pred_with_gt(pred_name, gt_csv)
    merge_csv.to_csv(osp.join(test_result_dir, 'pred_with_gt_infor.csv'))

    # tensorboard
    if args.tensorboard:
        subprocess.run(["tensorboard", "--logdir", exp_dir, "--port", "6007"])
